from django.contrib import admin
from .models import Blog
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from django.core.files.temp import NamedTemporaryFile
import requests
from bs4 import BeautifulSoup
from html import unescape


@admin.register(Blog)
class BlogAdmin(admin.ModelAdmin):
    list_display = ('title', 'published_date', 'image')
    actions = ['export_blogs_word']

    def export_blogs_word(self, request, queryset):
        # Sadece superuser kontrolü
        if not request.user.is_superuser:
            self.message_user(request, "Bu işlemi sadece admin yapabilir!", level='error')
            return None

        document = Document()
        document.add_heading('Blogs Report', 0)

        for blog in queryset.order_by('-published_date'):
            document.add_heading(blog.title, level=1)
            document.add_paragraph(
                f"Published: {blog.published_date.strftime('%Y-%m-%d %H:%M')}"
            )

            # HTML içeriği temizle ve ekle
            content = self.clean_html_content(blog.content)
            document.add_paragraph(content)

            # Cloudinary resim ekleme
            if blog.image:
                try:
                    img_url = blog.image.url

                    # Resmi indir
                    response = requests.get(img_url, timeout=10)
                    if response.status_code == 200:
                        img_temp = NamedTemporaryFile(delete=False, suffix='.jpg')
                        img_temp.write(response.content)
                        img_temp.flush()
                        img_temp.close()

                        # Word'e ekle
                        document.add_picture(img_temp.name, width=Inches(5))

                        # Geçici dosyayı sil
                        import os
                        os.unlink(img_temp.name)
                    else:
                        document.add_paragraph(f"Resim yüklenemedi (HTTP {response.status_code})")

                except Exception as e:
                    document.add_paragraph(f"Resim eklenemedi: {str(e)}")
            else:
                document.add_paragraph("Resim yok.")

            document.add_paragraph('\n' + '-' * 50 + '\n')

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="blogs.docx"'
        document.save(response)
        return response

    def clean_html_content(self, html_content):
        """HTML etiketlerini temizle ve düz metin döndür"""
        try:
            # BeautifulSoup ile HTML'i parse et
            soup = BeautifulSoup(html_content, 'html.parser')

            # Tüm metni al
            text = soup.get_text(separator='\n')

            # HTML entities'leri düzelt
            text = unescape(text)

            # Fazla boşlukları temizle
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return '\n'.join(lines)
        except:
            # Hata durumunda basit temizlik
            return html_content.replace('<p>', '').replace('</p>', '\n').replace('<br>', '\n')

    export_blogs_word.short_description = "Seçili blogları Word olarak indir"