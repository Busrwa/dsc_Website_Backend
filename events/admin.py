#event/admin.py
from django.contrib import admin
from .models import Event
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches
from django.core.files.temp import NamedTemporaryFile
import requests
from bs4 import BeautifulSoup
from html import unescape


@admin.register(Event)
class EventAdmin(admin.ModelAdmin):
    list_display = ('name', 'date', 'location', 'image')
    list_filter = ('date',)
    search_fields = ('name', 'location')
    actions = ['export_events_word']

    def export_events_word(self, request, queryset):
        if not request.user.is_superuser:
            self.message_user(request, "Bu işlemi sadece admin yapabilir!", level='error')
            return None

        document = Document()
        document.add_heading('Events Report', 0)

        for event in queryset.order_by('-date'):
            document.add_heading(event.name, level=1)
            document.add_paragraph(f"Date: {event.date.strftime('%Y-%m-%d %H:%M')}")
            document.add_paragraph(f"Location: {event.location}")

            # HTML içeriği temizle ve ekle
            description = self.clean_html_content(event.description)
            document.add_paragraph(description)

            # Cloudinary resim ekleme
            if event.image:
                try:
                    img_url = event.image.url

                    # Resmi indir
                    response = requests.get(img_url, timeout=10)
                    if response.status_code == 200:
                        img_temp = NamedTemporaryFile(delete=False, suffix='.jpg')
                        img_temp.write(response.content)
                        img_temp.flush()
                        img_temp.close()

                        # Word'e ekle
                        document.add_picture(img_temp.name, width=Inches(5))

                        # Geçici dosyayı sil
                        import os
                        os.unlink(img_temp.name)
                    else:
                        document.add_paragraph(f"Resim yüklenemedi (HTTP {response.status_code})")

                except Exception as e:
                    document.add_paragraph(f"Resim eklenemedi: {str(e)}")
            else:
                document.add_paragraph("Resim yok.")

            document.add_paragraph('\n' + '-' * 50 + '\n')

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="events.docx"'
        document.save(response)
        return response

    def clean_html_content(self, html_content):
        """HTML etiketlerini temizle ve düz metin döndür"""
        try:
            # BeautifulSoup ile HTML'i parse et
            soup = BeautifulSoup(html_content, 'html.parser')

            # Tüm metni al
            text = soup.get_text(separator='\n')

            # HTML entities'leri düzelt
            text = unescape(text)

            # Fazla boşlukları temizle
            lines = [line.strip() for line in text.split('\n') if line.strip()]
            return '\n'.join(lines)
        except:
            # Hata durumunda basit temizlik
            return html_content.replace('<p>', '').replace('</p>', '\n').replace('<br>', '\n')

    export_events_word.short_description = "Seçili eventleri Word olarak indir"